
import io
import re
import logging
from dataclasses import dataclass, field
from typing import Optional

import pdfplumber
import PyPDF2   
from docx import Document

logger = logging.getLogger(__name__)

@dataclass
class ParsedDocument:
    file_id: str
    file_name: str
    source: str = "gdrive"
    raw_text: str = ""
    page_count: int = 0
    mime_type: str = ""
    error: Optional[str] = None

    @property
    def is_valid(self) -> bool:
        return bool(self.raw_text.strip()) and self.error is None


class DocumentParser:
    """Extracts and normalizes text from various document types."""

    def parse(self, content: bytes, file_name: str, file_id: str, mime_type: str) -> ParsedDocument:
        doc = ParsedDocument(file_id=file_id, file_name=file_name, mime_type=mime_type)
        ext = file_name.lower().split(".")[-1] if "." in file_name else ""

        try:
            # Route by extension or MIME type
            if ext == "pdf" or mime_type == "application/pdf":
                doc.raw_text, doc.page_count = self._parse_pdf(content)
            elif ext in ("docx",) or "wordprocessingml" in mime_type:
                doc.raw_text = self._parse_docx(content)
            elif ext == "txt" or mime_type == "text/plain":
                doc.raw_text = self._parse_txt(content)
            else:
                # Fallback: try as text
                doc.raw_text = content.decode("utf-8", errors="replace")

            doc.raw_text = self._normalize(doc.raw_text)
            logger.info(f"Parsed {file_name}: {len(doc.raw_text)} chars, {doc.page_count} pages")

        except Exception as e:
            doc.error = str(e)
            logger.error(f"Failed to parse {file_name}: {e}")

        return doc

    # ------------------------------------------------------------------ #
    #  Parsers                                                             #
    # ------------------------------------------------------------------ #

    def _parse_pdf(self, content: bytes) -> tuple[str, int]:
        try:
            text_parts = []
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            return "\n\n".join(text_parts), len(text_parts)
        except Exception:
            # Fallback to PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
            return "\n\n".join(parts), len(parts)

    def _parse_docx(self, content: bytes) -> str:
        
        doc = Document(io.BytesIO(content))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        # Also extract table cells
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs)

    def _parse_txt(self, content: bytes) -> str:
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                return content.decode(enc)
            except UnicodeDecodeError:
                continue
        return content.decode("utf-8", errors="replace")

    # ------------------------------------------------------------------ #
    #  Normalization                                                       #
    # ------------------------------------------------------------------ #

    def _normalize(self, text: str) -> str:
        # Remove null bytes and non-printable chars
        text = text.replace("\x00", "")
        # Normalize whitespace
        text = re.sub(r"\r\n", "\n", text)
        text = re.sub(r"\r", "\n", text)
        # Collapse 3+ blank lines → 2
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Strip trailing whitespace per line
        lines = [line.rstrip() for line in text.split("\n")]
        text = "\n".join(lines)
        return text.strip()